#!/usr/bin/env python3
"""
Generates the free contractor invoice templates: a real PDF and a real editable
.docx, generic plus a variant for each trade.

    python3 scripts/build_templates.py

These are the backlink magnets. They only work if they are genuinely good — a
"free template" that turns out to be a lead-capture form with a watermark earns
nothing but a bounce. So: no email gate, no watermark, no signup. Download the
file, use it, keep it. If someone never installs the app but bookmarks the
template, that is still a link and still a person who now knows the name.

ON THE "GOOGLE DOCS VERSION". We cannot create a Google Doc from a script — it
needs a Google account and a shared "make a copy" link, which is a human action.
Rather than fake it with a dead link, we ship a real .docx (which Google Docs
opens and converts natively, and which Word and Pages open too) and the Google
Docs copy-link is listed as a manual task in seo-manual-tasks.md. When it exists,
put it in seo_config.GOOGLE_DOCS_TEMPLATE_URL and the page will link it
automatically.
"""
import os
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer, Table,
                                TableStyle)

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
sys.path.insert(0, HERE)
sys.path.insert(0, os.path.join(HERE, "content"))

from trades import TRADES  # noqa: E402

OUT = os.path.join(ROOT, "templates", "downloads")
ORANGE = colors.HexColor("#E8732C")
INK = colors.HexColor("#1a1d23")
MUTED = colors.HexColor("#6b7280")

# The generic starter rows. Every trade overrides these with its own.
GENERIC_ROWS = [
    ("Labour", "hrs", "", "", ""),
    ("Materials", "", "", "", ""),
    ("Call-out / service fee", "1", "", "", ""),
    ("", "", "", "", ""),
    ("", "", "", "", ""),
    ("", "", "", "", ""),
]


def trade_rows(slug):
    """Pre-fill the line-item column with the things that trade actually bills for,
    quantities and prices left blank. That is what makes a trade variant worth
    downloading rather than a find-and-replace on the title."""
    t = TRADES[slug]
    rows = [(item, "", "", "", "") for item, _unit, _rng in t["line_items"][:6]]
    while len(rows) < 6:
        rows.append(("", "", "", "", ""))
    return rows


# ------------------------------------------------------------------------ PDF --
def build_pdf(path, title, rows, note):
    doc = SimpleDocTemplate(path, pagesize=LETTER,
                            leftMargin=0.7 * inch, rightMargin=0.7 * inch,
                            topMargin=0.7 * inch, bottomMargin=0.7 * inch,
                            title=title, author="Toolbelt (toolbelt.pro)")
    ss = getSampleStyleSheet()
    h = ParagraphStyle("h", parent=ss["Title"], fontSize=26, textColor=INK,
                       alignment=0, spaceAfter=2, fontName="Helvetica-Bold")
    sub = ParagraphStyle("sub", parent=ss["Normal"], fontSize=9.5, textColor=MUTED)
    lbl = ParagraphStyle("lbl", parent=ss["Normal"], fontSize=8, textColor=MUTED,
                         fontName="Helvetica-Bold")
    val = ParagraphStyle("val", parent=ss["Normal"], fontSize=10.5, textColor=INK)
    small = ParagraphStyle("small", parent=ss["Normal"], fontSize=8, textColor=MUTED,
                           leading=11)

    E = "_" * 34
    story = []

    story.append(Paragraph("INVOICE", h))
    story.append(Paragraph(title, sub))
    story.append(Spacer(1, 16))

    meta = Table([
        [Paragraph("FROM", lbl), Paragraph("BILL TO", lbl)],
        [Paragraph(f"Business name{E}<br/>Address{E}<br/>{E}<br/>"
                   f"Phone{E}<br/>Email{E}<br/>Licence no.{E}", val),
         Paragraph(f"Client name{E}<br/>Job address{E}<br/>{E}<br/>"
                   f"Phone{E}<br/>Email{E}<br/>&nbsp;", val)],
    ], colWidths=[3.5 * inch, 3.5 * inch])
    meta.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 4),
        ("TOPPADDING", (0, 1), (-1, 1), 2),
    ]))
    story.append(meta)
    story.append(Spacer(1, 14))

    info = Table([[
        Paragraph(f"<b>Invoice #</b> {'_' * 14}", val),
        Paragraph(f"<b>Date</b> {'_' * 14}", val),
        Paragraph(f"<b>Due</b> {'_' * 14}", val),
    ]], colWidths=[2.33 * inch] * 3)
    info.setStyle(TableStyle([("BOTTOMPADDING", (0, 0), (-1, -1), 6)]))
    story.append(info)
    story.append(Spacer(1, 10))

    data = [["Description", "Qty", "Unit", "Rate", "Amount"]]
    for r in rows:
        data.append([Paragraph(r[0], val), r[1], r[2], r[3], r[4]])
    for _ in range(4):
        data.append(["", "", "", "", ""])
    data += [
        ["", "", "", "Subtotal", ""],
        ["", "", "", "Tax", ""],
        ["", "", "", "Deposit paid", ""],
        ["", "", "", "TOTAL DUE", ""],
    ]
    tbl = Table(data, colWidths=[3.5 * inch, 0.65 * inch, 0.7 * inch,
                                 1.0 * inch, 1.15 * inch])
    n = len(data)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), INK),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
        ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
        ("GRID", (0, 0), (-1, n - 5), 0.5, colors.HexColor("#dfe3e8")),
        ("LINEABOVE", (3, n - 4), (-1, n - 4), 0.8, colors.HexColor("#dfe3e8")),
        ("GRID", (3, n - 4), (-1, -1), 0.5, colors.HexColor("#dfe3e8")),
        ("FONTNAME", (3, n - 1), (-1, -1), "Helvetica-Bold"),
        ("BACKGROUND", (3, n - 1), (-1, -1), colors.HexColor("#fdf1e8")),
        ("TEXTCOLOR", (3, n - 1), (3, n - 1), ORANGE),
        ("TOPPADDING", (0, 1), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 1), (-1, -1), 7),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 14))

    story.append(Paragraph("<b>Payment terms</b>", lbl))
    story.append(Paragraph(
        f"Payment due within {E[:12]} days of the invoice date. "
        f"Accepted methods: {E[:22]}<br/>"
        f"Late payment: {E[:30]}", val))
    story.append(Spacer(1, 10))
    story.append(Paragraph("<b>Notes / scope of work</b>", lbl))
    story.append(Paragraph("<br/>".join(["_" * 92] * 3), val))
    story.append(Spacer(1, 18))

    story.append(Paragraph(note, small))
    story.append(Spacer(1, 6))
    story.append(Paragraph(
        "Free template from Toolbelt — the invoice app for contractors. "
        "toolbelt.pro. Use it, edit it, keep it. No attribution required.", small))
    doc.build(story)
    return path


# ----------------------------------------------------------------------- DOCX --
def build_docx(path, title, rows, note):
    d = Document()
    st = d.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)

    h = d.add_paragraph()
    r = h.add_run("INVOICE")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0x1A, 0x1D, 0x23)
    sub = d.add_paragraph()
    sr = sub.add_run(title)
    sr.font.size = Pt(9.5)
    sr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    meta = d.add_table(rows=2, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.cell(0, 0).text = "FROM"
    meta.cell(0, 1).text = "BILL TO"
    meta.cell(1, 0).text = ("Business name\nAddress\nPhone\nEmail\nLicence no.")
    meta.cell(1, 1).text = ("Client name\nJob address\nPhone\nEmail")
    d.add_paragraph()

    info = d.add_paragraph("Invoice #: ____________    Date: ____________    "
                           "Due: ____________")
    info.alignment = WD_ALIGN_PARAGRAPH.LEFT

    t = d.add_table(rows=1, cols=5)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, label in enumerate(["Description", "Qty", "Unit", "Rate", "Amount"]):
        hdr[i].text = label
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row in list(rows) + [("", "", "", "", "")] * 4:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    for label in ["Subtotal", "Tax", "Deposit paid", "TOTAL DUE"]:
        cells = t.add_row().cells
        cells[3].text = label
        if label == "TOTAL DUE":
            for p in cells[3].paragraphs:
                for run in p.runs:
                    run.bold = True

    d.add_paragraph()
    p = d.add_paragraph()
    p.add_run("Payment terms: ").bold = True
    p.add_run("Payment due within ______ days of the invoice date. "
              "Accepted methods: ______________________")
    p = d.add_paragraph()
    p.add_run("Notes / scope of work:").bold = True
    d.add_paragraph("____________________________________________________________")
    d.add_paragraph("____________________________________________________________")

    d.add_paragraph()
    f = d.add_paragraph()
    fr = f.add_run(note + "\n\nFree template from Toolbelt — the invoice app for "
                          "contractors. toolbelt.pro. Use it, edit it, keep it. "
                          "No attribution required.")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    d.save(path)
    return path


GENERIC_NOTE = ("Tip: the line items are the part that gets you paid without a phone call. "
                "\"Repair — $340\" invites a conversation. Three lines naming what you "
                "diagnosed, what you replaced and what you tested do not.")


def main():
    os.makedirs(OUT, exist_ok=True)
    made = []

    build_pdf(os.path.join(OUT, "contractor-invoice-template.pdf"),
              "Generic contractor invoice template", GENERIC_ROWS, GENERIC_NOTE)
    build_docx(os.path.join(OUT, "contractor-invoice-template.docx"),
               "Generic contractor invoice template", GENERIC_ROWS, GENERIC_NOTE)
    made.append("contractor-invoice-template")

    for slug, t in TRADES.items():
        note = (f"Tip for {t['name'].lower()}: {t['checklist'][0].lower()} — it is the "
                f"line most often missing from a {t['singular']}'s invoice, and the one "
                f"that most often causes the argument.")
        base = f"{slug}-invoice-template"
        build_pdf(os.path.join(OUT, base + ".pdf"),
                  f"{t['name']} invoice template", trade_rows(slug), note)
        build_docx(os.path.join(OUT, base + ".docx"),
                   f"{t['name']} invoice template", trade_rows(slug), note)
        made.append(base)

    for m in made:
        pdf = os.path.getsize(os.path.join(OUT, m + ".pdf")) / 1024
        docx = os.path.getsize(os.path.join(OUT, m + ".docx")) / 1024
        print(f"  {m:<38} pdf {pdf:>5.0f}KB   docx {docx:>5.0f}KB")
    print(f"\n{len(made)} template(s) x 2 formats -> templates/downloads/")


if __name__ == "__main__":
    main()
